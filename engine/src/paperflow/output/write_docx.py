"""DOCX export — rendered FROM the canonical ManuscriptState (not the source of record).

Uses python-docx. If python-docx is unavailable, returns False (like the best-effort PDF
compile) so the rest of the run still produces manuscript.json + HTML preview.
"""
from __future__ import annotations

import re
from pathlib import Path

from .manuscript_state import ManuscriptState

_HAS_DOCX = True
try:  # optional dependency
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:  # pragma: no cover - environment without python-docx
    _HAS_DOCX = False


def _bordered_box(doc):
    """A 1x1 bordered table used as a figure-placeholder box. Returns the cell."""
    t = doc.add_table(rows=1, cols=1)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "888888")
        borders.append(el)
    t._tbl.tblPr.append(borders)
    return t.cell(0, 0)


def available() -> bool:
    return _HAS_DOCX


_BOLD = re.compile(r"\*\*(.+?)\*\*")


def _add_inline(paragraph, text: str) -> None:
    """Render **bold** markers as real runs; everything else as plain runs."""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_markdown(doc, md: str) -> None:
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        h = re.match(r"^(#{1,4})\s+(.*)", line)
        li = re.match(r"^[-*]\s+(.*)", line)
        if h:
            doc.add_heading(h.group(2).strip(), level=min(4, len(h.group(1)) + 1))
        elif li:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, li.group(1).strip())
        else:
            _add_inline(doc.add_paragraph(), line)


def write_docx(ms: ManuscriptState, path: Path) -> bool:
    """Write the manuscript to a .docx file. Returns True on success, False if unavailable."""
    if not _HAS_DOCX:
        return False
    doc = Document()
    try:
        doc.styles["Normal"].font.size = Pt(11)
    except Exception:
        pass

    if ms.meta.title:
        doc.add_heading(ms.meta.title, level=0)
    if ms.meta.authors:
        doc.add_paragraph(ms.meta.authors).italic = True
    if ms.meta.affiliation:
        doc.add_paragraph(ms.meta.affiliation)
    if ms.meta.keywords:
        kw = doc.add_paragraph()
        kw.add_run("Keywords: ").bold = True
        kw.add_run(", ".join(ms.meta.keywords))

    for s in ms.sections:
        doc.add_heading(s.title, level=1)
        _add_markdown(doc, s.markdown)

    if ms.figures:
        doc.add_page_break()
        doc.add_heading("Figures & Tables — placeholders (to be created)", level=1)
        fig_n = tab_n = 0
        for f in ms.figures:
            if f.kind == "table":
                tab_n += 1
                label = f"Table {tab_n}"
            else:
                fig_n += 1
                label = f"Figure {fig_n}"
            # bordered box: WHAT to draw goes inside
            cell = _bordered_box(doc)
            head = cell.paragraphs[0]
            head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            head.add_run(f"[{label} — 여기에 이 그림을 그리세요]").bold = True
            body = cell.add_paragraph()
            body.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body.add_run(f.message)
            if f.generation_prompt:
                gp = cell.add_paragraph()
                gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = gp.add_run(f.generation_prompt[:220])
                r.italic = True
                r.font.size = Pt(9)
            # caption / title BELOW the box
            cap = doc.add_paragraph()
            cap.add_run(f"{label}. ").bold = True
            cap.add_run(f.caption_draft or f.message)
            if f.section:
                cap.add_run(f"  (→ {f.section})").italic = True
            doc.add_paragraph()  # spacer

    if ms.references:
        doc.add_heading("References", level=1)
        for r in ms.references:
            yr = f" ({r.year})" if r.year else ""
            doi = f" DOI: {r.doi}" if r.doi else ""
            doc.add_paragraph(f"[{r.key}] {r.authors}{yr}. {r.title}.{doi}", style="List Bullet")

    if ms.grounding.claims:
        doc.add_heading("Grounding report", level=1)
        s = ms.grounding.summary
        doc.add_paragraph(
            f"grounded: {s['grounded']} · partial: {s['partial']} · missing: {s['missing']}")
        for w in ms.grounding.warnings:
            doc.add_paragraph(f"⚠ {w}", style="List Bullet")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return True

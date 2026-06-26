"""Steps 6-7 — canonical manuscript state, HTML preview, DOCX export, figure placeholders."""
from pathlib import Path

import docx

from paperflow.figures import plan as figure_plan
from paperflow.output import manuscript_state as ms_mod
from paperflow.output import write_docx
from paperflow.schemas.claim import ClaimGraph, GEdge, GNode


def _graph_with_artifact() -> ClaimGraph:
    nodes = [
        GNode(id="C1", kind="claim", text="Material dominates rupture risk"),
        GNode(id="E1", kind="evidence", text="Material group ST=0.72"),
        GNode(id="D1", kind="data", text="sobol.csv", provenance="data/sobol/x.csv"),
        GNode(id="W1", kind="warrant", text="stress causes rupture"),
        GNode(id="F1", kind="artifact", text="Sobol bar chart", represented_as="figure",
              section="result"),
        GNode(id="C2", kind="claim", text="bare claim"),
    ]
    edges = [
        GEdge(src="E1", dst="C1", rel="supports"),
        GEdge(src="E1", dst="D1", rel="derived_from"),
        GEdge(src="W1", dst="C1", rel="justifies"),
        GEdge(src="F1", dst="E1", rel="visualizes"),
    ]
    return ClaimGraph(nodes=nodes, edges=edges)


def test_figure_plan_emits_caption_and_prompt():
    spec = figure_plan.plan("/nonexistent", _graph_with_artifact())
    figs = spec["figures"]
    assert figs and figs[0]["id"] == "F1"
    assert figs[0]["caption_draft"]
    assert "publication-quality" in figs[0]["generation_prompt"]
    assert "do not fabricate" in figs[0]["generation_prompt"]


def test_build_grounding_status():
    rep = ms_mod.build_grounding(_graph_with_artifact())
    by_id = {c.id: c for c in rep.claims}
    assert by_id["C1"].status == "grounded"       # evidence + data + warrant
    assert by_id["C2"].status == "missing"        # nothing
    assert rep.summary["grounded"] == 1 and rep.summary["missing"] == 1


def test_build_manuscript_state():
    sections = {"method": "## 2.1 Setup\nWe did **X**.", "result": "Material dominates.",
                "abstract": "Short abstract."}
    spec = figure_plan.plan("/nonexistent", _graph_with_artifact())
    refs = [{"key": "loree1992", "authors": "Loree", "year": 1992, "title": "Stress and rupture"}]
    ms = ms_mod.build(sections, graph=_graph_with_artifact(), figure_spec=spec,
                      meta={"title": "My Paper", "authors": "A B", "keywords": ["rupture"]},
                      found_references=refs, warnings=["claim_without_evidence: bare claim"])
    # abstract sorts before method/result in canonical order
    assert [s.name for s in ms.sections] == ["abstract", "method", "result"]
    assert ms.meta.title == "My Paper"
    assert ms.figures[0].id == "F1"
    assert ms.references[0].key == "loree1992"
    assert ms.grounding.warnings == ["claim_without_evidence: bare claim"]


def test_to_html_is_escaped_and_self_contained():
    sections = {"result": "Risk < 5% and a <script>alert(1)</script> tag."}
    ms = ms_mod.build(sections, graph=_graph_with_artifact(),
                      meta={"title": "T <b>x</b>"})
    h = ms_mod.to_html(ms)
    assert "<style>" in h and "pf-preview" in h
    # raw user/script content is escaped, not injected
    assert "<script>alert(1)</script>" not in h
    assert "&lt;script&gt;" in h
    assert "T &lt;b&gt;x&lt;/b&gt;" in h


def test_to_markdown_roundtrips_sections():
    ms = ms_mod.build({"method": "Did X.", "result": "Found Y."},
                      meta={"title": "P"})
    md = ms_mod.to_markdown(ms)
    assert "# P" in md and "## Methods" in md and "## Results" in md


def test_write_docx(tmp_path):
    assert write_docx.available()
    sections = {"method": "## Setup\nWe did **X**.\n- step one\n- step two",
                "result": "Material dominates."}
    spec = figure_plan.plan("/nonexistent", _graph_with_artifact())
    ms = ms_mod.build(sections, graph=_graph_with_artifact(), figure_spec=spec,
                      meta={"title": "My Paper", "authors": "A B"})
    out = tmp_path / "paper.docx"
    assert write_docx.write_docx(ms, out) is True
    assert out.is_file() and out.stat().st_size > 0
    # re-open and confirm content landed
    d = docx.Document(str(out))
    texts = "\n".join(p.text for p in d.paragraphs)
    assert "My Paper" in texts
    assert "Material dominates." in texts
    assert "step one" in texts
    assert any("F1" in p.text for p in d.paragraphs)   # figure placeholder present
